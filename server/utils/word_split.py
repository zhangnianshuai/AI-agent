"""
题库 Word 文档解析与分块工具（状态机版本）

Word 模板规范
=============
每道题按顺序包含 4 个字段，以 -字段名：值 格式书写：

    -题目：MySQL如何定位和优化慢查询？
    -评分标准：定位方法准确30分，优化方案合理40分，实际案例支撑30分
    -参考回复：通过slow_query_log开启慢查询日志，用EXPLAIN查看执行计划...
    -难度：4

    -题目：下一题...

字段别名兼容
============
题目  ← 问题、question
评分标准 ← 评分要点、打分标准、scoring、scoring_criteria
答案  ← 参考答案、参考回复、标准答案、答案、answer
难度  ← 难度等级、difficulty

状态机逻辑
==========
- 读到"题目"时，上一题必须已完成（状态=答案或难度），否则丢弃残缺题
- 缺失题目/评分标准/参考回复 → 丢弃整题
- 缺失难度（读到下一题"题目"时状态为"答案"）→ 默认设为 3
- 非字段行（续行）追加到当前字段值末尾

返回格式
========
list[dict], 每个 dict：
    {"question": str, "answer": str, "scoring_criteria": str | None, "difficulty": int}
"""

import re
from enum import Enum, auto
from pathlib import Path
from typing import Optional

from docx import Document


class FieldCategory(Enum):
    QUESTION = auto()
    SCORING = auto()
    ANSWER = auto()
    DIFFICULTY = auto()


# 字段名 → (类别, 输出 key)
_FIELD_MAP: dict[str, tuple[FieldCategory, str]] = {
    # ── question ──
    "题目":     (FieldCategory.QUESTION,   "question"),
    "问题":     (FieldCategory.QUESTION,   "question"),
    "question": (FieldCategory.QUESTION,   "question"),
    # ── scoring_criteria ──
    "评分标准":       (FieldCategory.SCORING, "scoring_criteria"),
    "评分要点":       (FieldCategory.SCORING, "scoring_criteria"),
    "打分标准":       (FieldCategory.SCORING, "scoring_criteria"),
    "scoring":        (FieldCategory.SCORING, "scoring_criteria"),
    "scoring_criteria": (FieldCategory.SCORING, "scoring_criteria"),
    # ── answer ──
    "答案":     (FieldCategory.ANSWER, "answer"),
    "参考答案": (FieldCategory.ANSWER, "answer"),
    "参考回复": (FieldCategory.ANSWER, "answer"),
    "标准答案": (FieldCategory.ANSWER, "answer"),
    "answer":   (FieldCategory.ANSWER, "answer"),
    # ── difficulty ──
    "难度":       (FieldCategory.DIFFICULTY, "difficulty"),
    "难度等级":   (FieldCategory.DIFFICULTY, "difficulty"),
    "difficulty": (FieldCategory.DIFFICULTY, "difficulty"),
}

# 有明确前缀的字段行：- 或 数字.
_FIELD_LINE_RE = re.compile(r"^(?:[-•]\s*|\d+\.\s*)(.+?)[：:](.*)$")
# 无前缀但行首直接是已知字段名（仅在状态机合法时才采纳）
_BARE_FIELD_RE = re.compile(r"^(.+?)[：:](.*)$")
# 编号标题行：13.xxx（没有：），可能是题目，仅在状态机允许时才当作题目
_NUMBERED_TITLE_RE = re.compile(r"^\d+\.\s*(.+)$")


class _ParseState(Enum):
    """解析状态机状态"""
    INIT = auto()    # 等待题目
    GOT_Q = auto()   # 已有题目，等待评分标准
    GOT_S = auto()   # 已有评分标准，等待答案
    GOT_A = auto()   # 已有答案，等待难度
    GOT_D = auto()   # 已有难度，题目完整


def _identify_field(raw_key: str) -> Optional[tuple[FieldCategory, str]]:
    """识别字段名，返回 (类别, 输出key)；无法识别返回 None"""
    return _FIELD_MAP.get(raw_key.strip())

def _emit(entry: dict, state: _ParseState) -> Optional[dict]:
    """
    根据状态机决定是否输出当前题目。
    返回 dict 表示输出，返回 None 表示丢弃。
    """
    if state == _ParseState.GOT_D:
        # 完整题目
        return entry
    if state == _ParseState.GOT_A:
        # 缺失难度，默认 3
        entry["difficulty"] = 3
        return entry
    # INIT / GOT_Q / GOT_S → 缺失关键字段，丢弃
    return None

# ═══════════════════════════════════════════════════════════════
# 段落式解析（状态机）
# ═══════════════════════════════════════════════════════════════

def _split_paragraphs(doc: Document) -> list[dict]:
    results: list[dict] = []
    state = _ParseState.INIT
    entry: dict = {"question": "", "answer": "", "scoring_criteria": None, "difficulty": 0}
    current_out_key: Optional[str] = None

    for p in doc.paragraphs:
        # 按 \n 拆行：兼容 Word 软回车（Shift+Enter）把多道题塞进同一个段落
        for raw_line in p.text.split('\n'):
            text = raw_line.strip()
            if not text:
                continue

            # 先尝试前缀匹配（- 或 数字.），再尝试无前缀匹配
            match = _FIELD_LINE_RE.match(text)
            has_prefix = True
            if not match:
                match = _BARE_FIELD_RE.match(text)
                has_prefix = False

            if not match:
                # 检查是否为编号标题（13.xxx）：状态机期待题目时，当题目处理
                title_match = _NUMBERED_TITLE_RE.match(text)
                if title_match and state in (_ParseState.INIT, _ParseState.GOT_D, _ParseState.GOT_A):
                    emitted = _emit(entry, state)
                    if emitted is not None:
                        results.append(emitted)
                    entry = {"question": title_match.group(1).strip(), "answer": "", "scoring_criteria": None, "difficulty": 0}
                    state = _ParseState.GOT_Q
                    current_out_key = "question"
                    continue

                # 完全不是字段行 → 续行
                if current_out_key and entry.get(current_out_key):
                    entry[current_out_key] += "\n" + text
                elif current_out_key:
                    entry[current_out_key] = text
                continue

            raw_key = match.group(1).strip()
            value = match.group(2).strip()
            identified = _identify_field(raw_key)

            if identified is None:
                # 无法识别的字段 → 续行
                if current_out_key and entry.get(current_out_key):
                    entry[current_out_key] += "\n" + text
                elif current_out_key:
                    entry[current_out_key] = text
                continue

            category, out_key = identified

            # ── 状态转换校验 ──
            valid_transition = _is_valid_transition(state, category)

            if not valid_transition and not has_prefix:
                # 无前缀 + 状态不合法 → 可能是答案正文，当续行
                if current_out_key and entry.get(current_out_key):
                    entry[current_out_key] += "\n" + text
                elif current_out_key:
                    entry[current_out_key] = text
                continue

            if not valid_transition:
                # 有前缀但状态不合法 → 残缺，丢弃
                entry = {"question": "", "answer": "", "scoring_criteria": None, "difficulty": 0}
                state = _ParseState.INIT
                current_out_key = None
                # 如果新字段是题目，可以开始新题
                if category == FieldCategory.QUESTION:
                    entry["question"] = value
                    state = _ParseState.GOT_Q
                    current_out_key = out_key
                continue

            # ── 合法转换 ──
            if category == FieldCategory.QUESTION:
                # 结算上一题
                emitted = _emit(entry, state)
                if emitted is not None:
                    results.append(emitted)
                entry = {"question": value, "answer": "", "scoring_criteria": None, "difficulty": 0}
                state = _ParseState.GOT_Q
                current_out_key = out_key

            elif category == FieldCategory.SCORING:
                entry[out_key] = value
                state = _ParseState.GOT_S
                current_out_key = out_key

            elif category == FieldCategory.ANSWER:
                entry[out_key] = value
                state = _ParseState.GOT_A
                current_out_key = out_key

            elif category == FieldCategory.DIFFICULTY:
                try:
                    entry[out_key] = int(value) if value else 3
                except ValueError:
                    entry[out_key] = 3
                state = _ParseState.GOT_D
                current_out_key = out_key

    # 文档末尾：结算最后一题
    emitted = _emit(entry, state)
    if emitted is not None:
        results.append(emitted)

    return results


def _is_valid_transition(state: _ParseState, category: FieldCategory) -> bool:
    """判断当前状态是否允许进入该字段类别"""
    _valid: dict[_ParseState, FieldCategory] = {
        _ParseState.INIT:  FieldCategory.QUESTION,
        _ParseState.GOT_Q: FieldCategory.SCORING,
        _ParseState.GOT_S: FieldCategory.ANSWER,
        _ParseState.GOT_A: FieldCategory.DIFFICULTY,
        _ParseState.GOT_D: FieldCategory.QUESTION,  # 完整题后可接下一题
        # GOT_A 遇到 QUESTION 也合法：缺难度时默认 3，由 _emit 处理
    }
    # GOT_A 遇 QUESTION 是合法的新题开始（难度默认3）
    if state == _ParseState.GOT_A and category == FieldCategory.QUESTION:
        return True
    return _valid.get(state) == category


# ═══════════════════════════════════════════════════════════════
# 表格格式解析（表头 → 按列映射）
# ═══════════════════════════════════════════════════════════════

def _split_tables(doc: Document) -> list[dict]:
    results: list[dict] = []

    for table in doc.tables:
        if not table.rows:
            continue

        # 表头
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        headers: list[Optional[tuple[FieldCategory, str]]] = [
            _identify_field(h) for h in header_cells
        ]

        # 数据行
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            entry: dict = {"question": "", "answer": "", "scoring_criteria": None, "difficulty": 0}

            for idx, identified in enumerate(headers):
                if identified is None or idx >= len(cells):
                    continue
                _, out_key = identified
                value = cells[idx]
                if not value:
                    continue
                if out_key == "difficulty":
                    try:
                        entry[out_key] = int(value)
                    except ValueError:
                        entry[out_key] = 0
                else:
                    entry[out_key] = value

            if entry["question"] and entry["answer"]:
                # 表格中缺评分标准可接受，缺难度默认3
                if not entry.get("scoring_criteria"):
                    entry["scoring_criteria"] = None
                if not entry.get("difficulty"):
                    entry["difficulty"] = 3
                results.append(entry)

    return results


def split_word(file_path: str | Path) -> list[dict]:
    """
    读取 .docx 文件，状态机解析，返回标准化题目列表。

    Parameters
    ----------
    file_path : str or Path
        .docx 文件路径

    Returns
    -------
    list[dict]
        [{"question": str, "answer": str, "scoring_criteria": str|None, "difficulty": int}, ...]
    """
    doc = Document(str(file_path))

    # 优先表格格式
    if doc.tables:
        return _split_tables(doc)

    return _split_paragraphs(doc)

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
    else:
        # 默认使用 store 目录下第一个 docx 文件（仅开发调试用）
        from pathlib import Path
        rag_dir = Path(__file__).resolve().parent.parent / "store" / "job_rag"
        files = list(rag_dir.glob("*.docx"))
        if not files:
            print("用法: python word_split.py <file_path>")
            print("  或: 将 .docx 文件放入 store/job_rag/ 目录")
            sys.exit(1)
        file_path = str(files[0])
        print(f"自动选择: {file_path}\n")

    for item in split_word(file_path):
        print(item)